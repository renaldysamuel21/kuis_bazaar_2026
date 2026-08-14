import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Soal bazaar - Revisi Pertanyaan.docx"
OUTPUT = ROOT / "src" / "data" / "questions.json"

CHARACTER_GROUPS = {
    "mudah": list(range(8, 18)) + list(range(43, 53)),
    "sedang": list(range(19, 29)) + list(range(54, 64)),
    "sulit": list(range(30, 40)) + list(range(65, 75)),
}

TRUE_FALSE_GROUPS = {
    "mudah": list(range(78, 88)) + list(range(113, 123)),
    "sedang": list(range(89, 99)) + list(range(124, 134)),
    "sulit": list(range(100, 110)) + list(range(135, 145)),
}


def value_after(label: str, text: str) -> str | None:
    for line in text.splitlines():
        marker_index = line.find(label)
        if marker_index >= 0:
            return line[marker_index + len(label) :].strip()
    return None


def clean_terminal_period(value: str) -> str:
    return value[:-1] if value.endswith(".") else value


def parse_marker(text: str) -> str:
    first_line = text.splitlines()[0]
    return "materi" if "★" in first_line else "tambahan"


def parse_character(text: str, difficulty: str, index: int) -> dict:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    answer = value_after("Jawaban:", text)
    explanation = value_after("Penjelasan:", text)
    verse = value_after("Ayat:", text)
    if not answer or not explanation or not verse or len(lines) < 5:
        raise ValueError(f"Soal tokoh tidak lengkap: {text!r}")
    return {
        "id": f"character-{difficulty}-{index:02d}",
        "kind": "character",
        "difficulty": difficulty,
        "source": parse_marker(text),
        "question": lines[1],
        "answer": clean_terminal_period(answer),
        "explanation": explanation,
        "verse": clean_terminal_period(verse),
    }


def parse_true_false(text: str, difficulty: str, index: int) -> dict:
    statement = value_after("Pernyataan:", text)
    answer_text = value_after("Jawaban:", text)
    verse = value_after("Ayat:", text)
    if not statement or not answer_text or not verse:
        raise ValueError(f"Soal benar/salah tidak lengkap: {text!r}")
    answer = clean_terminal_period(answer_text).upper()
    if answer not in {"BENAR", "SALAH"}:
        raise ValueError(f"Kunci tidak dikenal: {answer!r}")
    item = {
        "id": f"true-false-{difficulty}-{index:02d}",
        "kind": "true-false",
        "difficulty": difficulty,
        "source": parse_marker(text),
        "statement": statement,
        "answer": answer == "BENAR",
        "verse": clean_terminal_period(verse),
    }
    correction = value_after("Perbaikan:", text)
    explanation = value_after("Penjelasan:", text)
    if correction:
        item["correction"] = correction
    if explanation:
        item["explanation"] = explanation
    return item


def main() -> None:
    document = Document(SOURCE)
    character_questions = []
    true_false_questions = []

    for difficulty, paragraph_indices in CHARACTER_GROUPS.items():
        for position, paragraph_index in enumerate(paragraph_indices, start=1):
            character_questions.append(
                parse_character(document.paragraphs[paragraph_index].text, difficulty, position)
            )

    for difficulty, paragraph_indices in TRUE_FALSE_GROUPS.items():
        for position, paragraph_index in enumerate(paragraph_indices, start=1):
            true_false_questions.append(
                parse_true_false(document.paragraphs[paragraph_index].text, difficulty, position)
            )

    data = {
        "characterQuestions": character_questions,
        "trueFalseQuestions": true_false_questions,
    }
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Wrote {len(character_questions)} character questions")
    print(f"Wrote {len(true_false_questions)} true/false questions")
    print(OUTPUT)


if __name__ == "__main__":
    main()
